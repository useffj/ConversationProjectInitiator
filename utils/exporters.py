"""
utils/exporters.py
──────────────────
Helpers for exporting project artefacts to JSON, DOCX, and PDF.
"""
from __future__ import annotations

import io
import json
import re
from datetime import datetime

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def export_project_json(session_state) -> str:
    """Return a JSON string of all project data captured in session state."""
    payload = {
        "export_timestamp": datetime.now().isoformat(),
        "user_answers": dict(session_state.get("user_answers", {})),
        "ai_feedback": dict(session_state.get("ai_feedback", {})),
        "charter_markdown": session_state.get("charter_markdown", ""),
        "knowledge_area_data": dict(session_state.get("knowledge_area_data", {})),
    }
    return json.dumps(payload, indent=2, ensure_ascii=False)


def build_full_report(session_state) -> str:
    """
    Assemble all generated artefacts (charter + every knowledge-area document)
    into a single Markdown document suitable for download.
    """
    now = datetime.now().strftime("%B %d, %Y")
    parts: list[str] = [
        "# Complete Project Documentation Package\n",
        f"*Generated: {now} — Conversational Project Initiator*\n\n",
        "---\n\n",
    ]

    charter = session_state.get("charter_markdown", "")
    if charter:
        parts.append(charter)
        parts.append("\n\n---\n\n")

    knowledge_data: dict = session_state.get("knowledge_area_data", {})
    for area, content in knowledge_data.items():
        parts.append(f"## {area}\n\n")
        parts.append(content)
        parts.append("\n\n---\n\n")

    return "".join(parts)


def _markdown_to_plain_text(markdown_text: str) -> str:
    """Convert markdown into readable plain text for DOCX/PDF export."""
    text = markdown_text
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\|", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    """Create a DOCX file in memory from markdown-like text."""
    doc = Document()
    plain = _markdown_to_plain_text(markdown_text)

    for line in plain.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def markdown_to_pdf_bytes(markdown_text: str) -> bytes:
    """Create a basic PDF file in memory from markdown-like text."""
    plain = _markdown_to_plain_text(markdown_text)
    buffer = io.BytesIO()

    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 50
    line_height = 14

    for raw_line in plain.splitlines():
        line = raw_line.strip()
        if not line:
            y -= line_height
            if y < 50:
                c.showPage()
                y = height - 50
            continue

        max_chars = 95
        chunks = [line[i:i + max_chars] for i in range(0, len(line), max_chars)]
        for chunk in chunks:
            c.drawString(40, y, chunk)
            y -= line_height
            if y < 50:
                c.showPage()
                y = height - 50

    c.save()
    buffer.seek(0)
    return buffer.getvalue()
